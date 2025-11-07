from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, UploadFile, File
from fastapi.security import OAuth2PasswordRequestForm
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
import io
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import pandas as pd
from auth import (
    get_password_hash,
    verify_password,
    create_access_token,
    decode_token,
    oauth2_scheme
)
from email_service import (
    generate_reset_token,
    get_token_expiry,
    send_password_reset_email,
    send_invite_email
)
from validators import Validators, ValidationError
from rate_limiter import (
    rate_limit_login,
    rate_limit_register,
    rate_limit_password_reset,
    rate_limit_invite,
    Request as RateLimitRequest
)

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
# Configure SSL/TLS for cloud database connections
if "mongodb+srv" in mongo_url:
    import ssl
    client = AsyncIOMotorClient(
        mongo_url,
        tls=True,
        tlsAllowInvalidCertificates=True,  # For Emergent environment
        serverSelectionTimeoutMS=10000,
        connectTimeoutMS=10000
    )
else:
    client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

app = FastAPI()
api_router = APIRouter(prefix="/api")


# ============ MODELS ============

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    username: str
    email: EmailStr
    hashed_password: str
    full_name: str
    is_active: bool = True
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class UserCreate(BaseModel):
    username: str
    email: EmailStr
    password: str
    full_name: str

class UserResponse(BaseModel):
    id: str
    username: str
    email: str
    full_name: str
    is_active: bool

class Token(BaseModel):
    access_token: str
    token_type: str
    user: UserResponse

class ProjectMember(BaseModel):
    user_id: str
    username: str
    role: str  # admin, editor, viewer
    added_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Project(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    description: Optional[str] = ""
    owner_id: str
    members: List[ProjectMember] = []
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ProjectCreate(BaseModel):
    name: str
    description: Optional[str] = ""

class AddMemberRequest(BaseModel):
    username: str
    role: str  # admin, editor, viewer

class Comment(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    text: str
    created_by: str
    user_id: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class TestCase(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_id: str
    tab: str = "General"
    title: str
    description: str
    priority: str
    type: str
    steps: str
    expected_result: str
    actual_result: Optional[str] = ""
    status: str = "draft"
    assigned_to: Optional[str] = ""
    executed_at: Optional[datetime] = None
    comments: List[Comment] = []
    is_template: bool = False
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class TestCaseCreate(BaseModel):
    project_id: str
    tab: Optional[str] = "General"
    title: str
    description: str
    priority: str
    type: str
    steps: str
    expected_result: str
    actual_result: Optional[str] = ""
    assigned_to: Optional[str] = ""
    is_template: Optional[bool] = False

class TestCaseUpdate(BaseModel):
    tab: Optional[str] = None
    title: Optional[str] = None
    description: Optional[str] = None
    priority: Optional[str] = None
    type: Optional[str] = None
    steps: Optional[str] = None
    expected_result: Optional[str] = None
    actual_result: Optional[str] = None
    status: Optional[str] = None
    assigned_to: Optional[str] = None
    executed_at: Optional[datetime] = None

class StatusUpdate(BaseModel):
    status: str

class BulkStatusUpdate(BaseModel):
    test_case_ids: List[str]
    status: str

class CommentCreate(BaseModel):
    test_case_id: str
    text: str

class ForgotPasswordRequest(BaseModel):
    email: EmailStr

class ResetPasswordRequest(BaseModel):
    token: str
    new_password: str



class InviteMemberRequest(BaseModel):
    email: EmailStr
    role: str

class AcceptInviteRequest(BaseModel):
    token: str


# ============ AUTH DEPENDENCIES ============

async def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    
    payload = decode_token(token)
    if payload is None:
        raise credentials_exception
    
    user_id: str = payload.get("sub")
    if user_id is None:
        raise credentials_exception
    
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if user is None:
        raise credentials_exception
    
    if not user.get("is_active", False):
        raise HTTPException(status_code=400, detail="Inactive user")
    
    return user

async def check_project_permission(project_id: str, user: dict, required_role: str = "viewer"):
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Owner has all permissions
    if project["owner_id"] == user["id"]:
        return project
    
    # Check if user is a member
    member = next((m for m in project.get("members", []) if m["user_id"] == user["id"]), None)
    if not member:
        raise HTTPException(status_code=403, detail="You don't have access to this project")
    
    # Check role permissions
    role_hierarchy = {"viewer": 0, "editor": 1, "admin": 2}
    if role_hierarchy.get(member["role"], 0) < role_hierarchy.get(required_role, 0):
        raise HTTPException(status_code=403, detail=f"You need {required_role} role for this action")
    
    return project


# ============ AUTH ROUTES ============

@api_router.post("/auth/register", response_model=Token)
async def register(user_data: UserCreate, request: RateLimitRequest = None):
    # Rate limiting
    if request:
        rate_limit_register(request)
    
    # Validate input
    try:
        username = Validators.validate_username(user_data.username)
        email = Validators.validate_email(user_data.email)
        full_name = Validators.validate_full_name(user_data.full_name)
        Validators.validate_password(user_data.password)
    except ValidationError as e:
        raise HTTPException(status_code=400, detail=str(e))
    
    # Check if username exists
    existing_user = await db.users.find_one({"username": username})
    if existing_user:
        raise HTTPException(status_code=400, detail="Username already registered")
    
    # Check if email exists
    existing_email = await db.users.find_one({"email": email})
    if existing_email:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create user
    hashed_password = get_password_hash(user_data.password)
    user = User(
        username=username,
        email=email,
        hashed_password=hashed_password,
        full_name=full_name
    )
    
    user_dict = user.model_dump()
    user_dict['created_at'] = user_dict['created_at'].isoformat()
    
    await db.users.insert_one(user_dict)
    
    # Create access token
    access_token = create_access_token(data={"sub": user.id})
    
    user_response = UserResponse(
        id=user.id,
        username=user.username,
        email=user.email,
        full_name=user.full_name,
        is_active=user.is_active
    )
    
    return Token(access_token=access_token, token_type="bearer", user=user_response)

@api_router.post("/auth/login", response_model=Token)
async def login(form_data: OAuth2PasswordRequestForm = Depends(), request: RateLimitRequest = None):
    # Rate limiting
    if request:
        rate_limit_login(request)
    
    # Validate input
    if not form_data.username or not form_data.password:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Username and password are required",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    user = await db.users.find_one({"username": form_data.username.strip()}, {"_id": 0})
    if not user or not verify_password(form_data.password, user["hashed_password"]):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    if not user.get("is_active", False):
        raise HTTPException(status_code=400, detail="Inactive user")
    
    access_token = create_access_token(data={"sub": user["id"]})
    
    user_response = UserResponse(
        id=user["id"],
        username=user["username"],
        email=user["email"],
        full_name=user["full_name"],
        is_active=user["is_active"]
    )
    
    return Token(access_token=access_token, token_type="bearer", user=user_response)

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(current_user: dict = Depends(get_current_user)):
    return UserResponse(
        id=current_user["id"],
        username=current_user["username"],
        email=current_user["email"],
        full_name=current_user["full_name"],
        is_active=current_user["is_active"]
    )

@api_router.post("/auth/forgot-password")
async def forgot_password(request: ForgotPasswordRequest):
    """Request password reset email"""
    user = await db.users.find_one({"email": request.email}, {"_id": 0})
    
    # Always return success to prevent email enumeration
    if not user:
        return {"message": "If an account with that email exists, a password reset link has been sent."}
    
    # Generate reset token
    reset_token = generate_reset_token()
    token_expiry = get_token_expiry()
    
    # Store reset token in database
    await db.users.update_one(
        {"email": request.email},
        {
            "$set": {
                "reset_token": reset_token,
                "reset_token_expiry": token_expiry.isoformat()
            }
        }
    )
    
    # Send email
    await send_password_reset_email(request.email, reset_token, user["username"])
    
    return {"message": "If an account with that email exists, a password reset link has been sent."}

@api_router.post("/auth/reset-password")
async def reset_password(request: ResetPasswordRequest):
    """Reset password using token"""
    # Find user with valid token
    user = await db.users.find_one({
        "reset_token": request.token
    }, {"_id": 0})
    
    if not user:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    
    # Check if token is expired
    token_expiry = datetime.fromisoformat(user["reset_token_expiry"])
    if datetime.now(timezone.utc) > token_expiry:
        raise HTTPException(status_code=400, detail="Reset token has expired")
    
    # Validate new password
    if len(request.new_password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    
    # Update password and clear reset token
    hashed_password = get_password_hash(request.new_password)
    await db.users.update_one(
        {"id": user["id"]},
        {
            "$set": {
                "hashed_password": hashed_password
            },
            "$unset": {
                "reset_token": "",
                "reset_token_expiry": ""
            }
        }
    )
    
    return {"message": "Password has been reset successfully. You can now login with your new password."}


# ============ PROJECT ROUTES ============

@api_router.post("/projects", response_model=Project)
async def create_project(input: ProjectCreate, current_user: dict = Depends(get_current_user)):
    project = Project(
        name=input.name,
        description=input.description,
        owner_id=current_user["id"],
        members=[]
    )
    
    doc = project.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.projects.insert_one(doc)
    return project

@api_router.get("/projects", response_model=List[Project])
async def get_projects(current_user: dict = Depends(get_current_user)):
    # Get projects where user is owner or member
    projects = await db.projects.find({
        "$or": [
            {"owner_id": current_user["id"]},
            {"members.user_id": current_user["id"]}
        ]
    }, {"_id": 0}).to_list(1000)
    
    for project in projects:
        if isinstance(project['created_at'], str):
            project['created_at'] = datetime.fromisoformat(project['created_at'])
    
    return projects

@api_router.get("/projects/{project_id}", response_model=Project)
async def get_project(project_id: str, current_user: dict = Depends(get_current_user)):
    project = await check_project_permission(project_id, current_user, "viewer")
    
    if isinstance(project['created_at'], str):
        project['created_at'] = datetime.fromisoformat(project['created_at'])
    
    return project

@api_router.delete("/projects/{project_id}")
async def delete_project(project_id: str, current_user: dict = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Only owner can delete
    if project["owner_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="Only project owner can delete the project")
    
    await db.test_cases.delete_many({"project_id": project_id})
    await db.projects.delete_one({"id": project_id})
    
    return {"message": "Project deleted successfully"}


@api_router.put("/projects/{project_id}")
async def rename_project(
    project_id: str,
    name: str,
    description: str = "",
    current_user: dict = Depends(get_current_user)
):
    project = await db.projects.find_one({"id": project_id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Only owner can rename
    if project["owner_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="Only project owner can rename the project")
    
    await db.projects.update_one(
        {"id": project_id},
        {"$set": {"name": name, "description": description}}
    )
    
    return {"message": "Project updated successfully"}


@api_router.post("/projects/{project_id}/members")
async def add_project_member(
    project_id: str,
    member_data: AddMemberRequest,
    current_user: dict = Depends(get_current_user)
):
    project = await check_project_permission(project_id, current_user, "admin")
    
    # Find user by username
    user = await db.users.find_one({"username": member_data.username}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Check if already a member
    if any(m["user_id"] == user["id"] for m in project.get("members", [])):
        raise HTTPException(status_code=400, detail="User is already a member")
    
    # Add member
    member = ProjectMember(
        user_id=user["id"],
        username=user["username"],
        role=member_data.role
    )
    
    member_dict = member.model_dump()
    member_dict['added_at'] = member_dict['added_at'].isoformat()
    
    await db.projects.update_one(
        {"id": project_id},
        {"$push": {"members": member_dict}}
    )
    
    return {"message": f"User {member_data.username} added as {member_data.role}"}

@api_router.put("/projects/{project_id}/members/{user_id}/role")
async def update_member_role(
    project_id: str,
    user_id: str,
    role: str,
    current_user: dict = Depends(get_current_user)
):
    project = await check_project_permission(project_id, current_user, "admin")
    
    await db.projects.update_one(
        {"id": project_id, "members.user_id": user_id},
        {"$set": {"members.$.role": role}}
    )
    
    return {"message": "Role updated successfully"}

@api_router.delete("/projects/{project_id}/members/{user_id}")
async def remove_project_member(
    project_id: str,
    user_id: str,
    current_user: dict = Depends(get_current_user)
):
    project = await check_project_permission(project_id, current_user, "admin")
    
    await db.projects.update_one(
        {"id": project_id},
        {"$pull": {"members": {"user_id": user_id}}}
    )
    
    return {"message": "Member removed successfully"}


# ============ PROJECT INVITE ENDPOINTS ============

@api_router.post("/projects/{project_id}/invites")
async def invite_member_by_email(
    project_id: str,
    invite_data: InviteMemberRequest,
    current_user: dict = Depends(get_current_user)
):
    """Send email invite to add a member to the project"""
    project = await check_project_permission(project_id, current_user, "admin")
    
    # Generate invite token
    invite_token = generate_reset_token()
    invite_expiry = get_token_expiry()  # 1 hour expiry
    
    # Store invite in database
    invite = {
        "id": str(uuid.uuid4()),
        "project_id": project_id,
        "project_name": project["name"],
        "email": invite_data.email,
        "role": invite_data.role,
        "token": invite_token,
        "expires_at": invite_expiry.isoformat(),
        "invited_by": current_user["id"],
        "invited_by_username": current_user["username"],
        "created_at": datetime.now(timezone.utc).isoformat(),
        "status": "pending"  # pending, accepted, expired
    }
    
    await db.invites.insert_one(invite)
    
    # Send email with invite link
    frontend_url = os.environ.get('FRONTEND_URL', 'https://testcenter.preview.emergentagent.com')
    invite_link = f"{frontend_url}/accept-invite/{invite_token}"
    
    # Send email using SendGrid
    await send_invite_email(
        invite_data.email,
        invite_token,
        project["name"],
        current_user["username"],
        invite_data.role
    )
    
    return {
        "message": f"Invitation sent to {invite_data.email}",
        "invite_link": invite_link  # For testing purposes
    }

@api_router.get("/invites/{token}")
async def get_invite_details(token: str):
    """Get invite details (doesn't require authentication)"""
    invite = await db.invites.find_one({"token": token}, {"_id": 0})
    
    if not invite:
        raise HTTPException(status_code=404, detail="Invite not found")
    
    # Check if expired
    if datetime.fromisoformat(invite["expires_at"]) < datetime.now(timezone.utc):
        raise HTTPException(status_code=400, detail="Invite has expired")
    
    if invite["status"] != "pending":
        raise HTTPException(status_code=400, detail="Invite already used")
    
    return {
        "project_name": invite["project_name"],
        "role": invite["role"],
        "invited_by": invite["invited_by_username"],
        "email": invite["email"]
    }

@api_router.post("/invites/{token}/accept")
async def accept_invite(
    token: str,
    current_user: dict = Depends(get_current_user)
):
    """Accept project invite (user must be logged in)"""
    invite = await db.invites.find_one({"token": token}, {"_id": 0})
    
    if not invite:
        raise HTTPException(status_code=404, detail="Invite not found")
    
    # Check if expired
    if datetime.fromisoformat(invite["expires_at"]) < datetime.now(timezone.utc):
        raise HTTPException(status_code=400, detail="Invite has expired")
    
    if invite["status"] != "pending":
        raise HTTPException(status_code=400, detail="Invite already used")
    
    # Verify email matches (optional - can be removed if you want any logged in user to accept)
    if current_user["email"] != invite["email"]:
        raise HTTPException(
            status_code=403,
            detail="This invite is for a different email address"
        )
    
    # Check if already a member
    project = await db.projects.find_one({"id": invite["project_id"]}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    if any(m["user_id"] == current_user["id"] for m in project.get("members", [])):
        raise HTTPException(status_code=400, detail="You are already a member of this project")
    
    # Add member to project
    member = ProjectMember(
        user_id=current_user["id"],
        username=current_user["username"],
        role=invite["role"]
    )
    
    member_dict = member.model_dump()
    member_dict['added_at'] = member_dict['added_at'].isoformat()
    
    await db.projects.update_one(
        {"id": invite["project_id"]},
        {"$push": {"members": member_dict}}
    )
    
    # Mark invite as accepted
    await db.invites.update_one(
        {"token": token},
        {"$set": {"status": "accepted", "accepted_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    return {
        "message": "Invite accepted successfully",
        "project_id": invite["project_id"],
        "project_name": invite["project_name"]
    }


@api_router.get("/projects/{project_id}/tabs")
async def get_project_tabs(project_id: str, current_user: dict = Depends(get_current_user)):
    await check_project_permission(project_id, current_user, "viewer")
    
    test_cases = await db.test_cases.find({"project_id": project_id}, {"_id": 0, "tab": 1}).to_list(1000)
    tabs = list(set([tc.get('tab', 'General') for tc in test_cases]))
    if not tabs:
        tabs = ["General"]
    return {"tabs": sorted(tabs)}

@api_router.get("/projects/{project_id}/statistics")
async def get_project_statistics(project_id: str, current_user: dict = Depends(get_current_user)):
    await check_project_permission(project_id, current_user, "viewer")
    
    test_cases = await db.test_cases.find({"project_id": project_id, "is_template": False}, {"_id": 0}).to_list(1000)
    
    stats = {
        "total": len(test_cases),
        "draft": len([tc for tc in test_cases if tc.get('status') == 'draft']),
        "success": len([tc for tc in test_cases if tc.get('status') == 'success']),
        "fail": len([tc for tc in test_cases if tc.get('status') == 'fail']),
        "by_priority": {
            "low": len([tc for tc in test_cases if tc.get('priority') == 'low']),
            "medium": len([tc for tc in test_cases if tc.get('priority') == 'medium']),
            "high": len([tc for tc in test_cases if tc.get('priority') == 'high'])
        },
        "by_tab": {}
    }
    
    for tc in test_cases:
        tab = tc.get('tab', 'General')
        if tab not in stats["by_tab"]:
            stats["by_tab"][tab] = {"total": 0, "draft": 0, "success": 0, "fail": 0}
        stats["by_tab"][tab]["total"] += 1
        stats["by_tab"][tab][tc.get('status', 'draft')] += 1
    
    return stats


# ============ TEST CASE ROUTES ============

@api_router.post("/test-cases", response_model=TestCase)
async def create_test_case(input: TestCaseCreate, current_user: dict = Depends(get_current_user)):
    await check_project_permission(input.project_id, current_user, "editor")
    
    test_case = TestCase(
        **input.model_dump(),
        created_by=current_user["username"]
    )
    
    doc = test_case.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    doc['updated_at'] = doc['updated_at'].isoformat()
    if doc.get('executed_at'):
        doc['executed_at'] = doc['executed_at'].isoformat()
    
    await db.test_cases.insert_one(doc)
    return test_case

@api_router.get("/test-cases", response_model=List[TestCase])
async def get_test_cases(
    project_id: Optional[str] = None,
    tab: Optional[str] = None,
    is_template: Optional[bool] = None,
    search: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    if project_id:
        await check_project_permission(project_id, current_user, "viewer")
    
    query = {}
    if project_id:
        query["project_id"] = project_id
    if tab:
        query["tab"] = tab
    if is_template is not None:
        query["is_template"] = is_template
    
    test_cases = await db.test_cases.find(query, {"_id": 0}).to_list(1000)
    
    for tc in test_cases:
        if isinstance(tc['created_at'], str):
            tc['created_at'] = datetime.fromisoformat(tc['created_at'])
        if isinstance(tc['updated_at'], str):
            tc['updated_at'] = datetime.fromisoformat(tc['updated_at'])
        if tc.get('executed_at') and isinstance(tc['executed_at'], str):
            tc['executed_at'] = datetime.fromisoformat(tc['executed_at'])
    
    if search:
        search_lower = search.lower()
        test_cases = [tc for tc in test_cases if 
                      search_lower in tc.get('title', '').lower() or 
                      search_lower in tc.get('description', '').lower() or
                      search_lower in tc.get('steps', '').lower()]
    
    return test_cases

@api_router.post("/test-cases/{test_case_id}/duplicate", response_model=TestCase)
async def duplicate_test_case(test_case_id: str, current_user: dict = Depends(get_current_user)):
    test_case = await db.test_cases.find_one({"id": test_case_id}, {"_id": 0})
    if not test_case:
        raise HTTPException(status_code=404, detail="Test case not found")
    
    await check_project_permission(test_case['project_id'], current_user, "editor")
    
    new_tc = TestCase(**test_case)
    new_tc.id = str(uuid.uuid4())
    new_tc.title = f"{test_case['title']} (Copy)"
    new_tc.created_at = datetime.now(timezone.utc)
    new_tc.updated_at = datetime.now(timezone.utc)
    new_tc.comments = []
    new_tc.created_by = current_user["username"]
    
    doc = new_tc.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    doc['updated_at'] = doc['updated_at'].isoformat()
    if doc.get('executed_at'):
        doc['executed_at'] = doc['executed_at'].isoformat()
    
    await db.test_cases.insert_one(doc)
    return new_tc

@api_router.put("/test-cases/{test_case_id}", response_model=TestCase)
async def update_test_case(
    test_case_id: str,
    input: TestCaseUpdate,
    current_user: dict = Depends(get_current_user)
):
    test_case = await db.test_cases.find_one({"id": test_case_id}, {"_id": 0})
    if not test_case:
        raise HTTPException(status_code=404, detail="Test case not found")
    
    await check_project_permission(test_case['project_id'], current_user, "editor")
    
    update_data = {k: v for k, v in input.model_dump().items() if v is not None}
    update_data['updated_at'] = datetime.now(timezone.utc).isoformat()
    
    if 'executed_at' in update_data and update_data['executed_at']:
        update_data['executed_at'] = update_data['executed_at'].isoformat()
    
    await db.test_cases.update_one(
        {"id": test_case_id},
        {"$set": update_data}
    )
    
    updated_test_case = await db.test_cases.find_one({"id": test_case_id}, {"_id": 0})
    
    if isinstance(updated_test_case['created_at'], str):
        updated_test_case['created_at'] = datetime.fromisoformat(updated_test_case['created_at'])
    if isinstance(updated_test_case['updated_at'], str):
        updated_test_case['updated_at'] = datetime.fromisoformat(updated_test_case['updated_at'])
    if updated_test_case.get('executed_at') and isinstance(updated_test_case['executed_at'], str):
        updated_test_case['executed_at'] = datetime.fromisoformat(updated_test_case['executed_at'])
    
    return updated_test_case

@api_router.patch("/test-cases/{test_case_id}/status", response_model=TestCase)
async def update_test_case_status(
    test_case_id: str,
    input: StatusUpdate,
    current_user: dict = Depends(get_current_user)
):
    test_case = await db.test_cases.find_one({"id": test_case_id}, {"_id": 0})
    if not test_case:
        raise HTTPException(status_code=404, detail="Test case not found")
    
    await check_project_permission(test_case['project_id'], current_user, "editor")
    
    update_data = {
        "status": input.status,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    if input.status in ['success', 'fail']:
        update_data['executed_at'] = datetime.now(timezone.utc).isoformat()
    
    await db.test_cases.update_one(
        {"id": test_case_id},
        {"$set": update_data}
    )
    
    updated_test_case = await db.test_cases.find_one({"id": test_case_id}, {"_id": 0})
    
    if isinstance(updated_test_case['created_at'], str):
        updated_test_case['created_at'] = datetime.fromisoformat(updated_test_case['created_at'])
    if isinstance(updated_test_case['updated_at'], str):
        updated_test_case['updated_at'] = datetime.fromisoformat(updated_test_case['updated_at'])
    if updated_test_case.get('executed_at') and isinstance(updated_test_case['executed_at'], str):
        updated_test_case['executed_at'] = datetime.fromisoformat(updated_test_case['executed_at'])
    
    return updated_test_case

@api_router.post("/test-cases/bulk-status")
async def bulk_update_status(input: BulkStatusUpdate, current_user: dict = Depends(get_current_user)):
    # Verify all test cases belong to projects user has access to
    test_cases = await db.test_cases.find({"id": {"$in": input.test_case_ids}}, {"_id": 0, "project_id": 1}).to_list(1000)
    
    for tc in test_cases:
        await check_project_permission(tc['project_id'], current_user, "editor")
    
    update_data = {
        "status": input.status,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    if input.status in ['success', 'fail']:
        update_data['executed_at'] = datetime.now(timezone.utc).isoformat()
    
    result = await db.test_cases.update_many(
        {"id": {"$in": input.test_case_ids}},
        {"$set": update_data}
    )
    
    return {"updated_count": result.modified_count}

@api_router.delete("/test-cases/bulk")
async def bulk_delete_test_cases(test_case_ids: List[str], current_user: dict = Depends(get_current_user)):
    test_cases = await db.test_cases.find({"id": {"$in": test_case_ids}}, {"_id": 0, "project_id": 1}).to_list(1000)
    
    for tc in test_cases:
        await check_project_permission(tc['project_id'], current_user, "editor")
    
    result = await db.test_cases.delete_many({"id": {"$in": test_case_ids}})
    return {"deleted_count": result.deleted_count}

@api_router.delete("/test-cases/{test_case_id}")
async def delete_test_case(test_case_id: str, current_user: dict = Depends(get_current_user)):
    test_case = await db.test_cases.find_one({"id": test_case_id}, {"_id": 0})
    if not test_case:
        raise HTTPException(status_code=404, detail="Test case not found")
    
    await check_project_permission(test_case['project_id'], current_user, "editor")
    
    await db.test_cases.delete_one({"id": test_case_id})
    return {"message": "Test case deleted successfully"}

@api_router.post("/test-cases/{test_case_id}/comments", response_model=Comment)
async def add_comment(
    test_case_id: str,
    input: CommentCreate,
    current_user: dict = Depends(get_current_user)
):
    test_case = await db.test_cases.find_one({"id": test_case_id})
    if not test_case:
        raise HTTPException(status_code=404, detail="Test case not found")
    
    await check_project_permission(test_case['project_id'], current_user, "viewer")
    
    comment = Comment(
        text=input.text,
        created_by=current_user["username"],
        user_id=current_user["id"]
    )
    comment_dict = comment.model_dump()
    comment_dict['created_at'] = comment_dict['created_at'].isoformat()
    
    await db.test_cases.update_one(
        {"id": test_case_id},
        {"$push": {"comments": comment_dict}}
    )
    
    return comment



# ============ IMPORT TEST CASES ============

@api_router.post("/test-cases/import/{project_id}")
async def import_test_cases(
    project_id: str,
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Import test cases from CSV or Excel file with validation"""
    # Check project permission
    project = await check_project_permission(project_id, current_user, "editor")
    
    # Validate file type
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    
    allowed_extensions = ['.csv', '.xlsx', '.xls']
    file_ext = os.path.splitext(file.filename)[1].lower()
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file format: {file_ext}. Allowed formats: CSV (.csv), Excel (.xlsx, .xls)"
        )
    
    # Read file contents
    try:
        contents = await file.read()
        
        # Parse file based on type
        if file_ext == '.csv':
            df = pd.read_csv(io.BytesIO(contents))
        else:  # xlsx or xls
            df = pd.read_excel(io.BytesIO(contents))
        
        # Validate required columns
        required_columns = ['Title', 'Description', 'Priority', 'Type', 'Steps', 'Expected Result']
        missing_columns = [col for col in required_columns if col not in df.columns and col.lower() not in df.columns]
        
        if missing_columns:
            # Create file format guide
            format_guide = """
Required CSV/Excel Format:
--------------------------
Columns (case-insensitive):
1. Title (required) - Test case title
2. Description (required) - Test case description  
3. Priority (required) - low, medium, or high
4. Type (required) - functional, negative, UI/UX, smoke, regression, or API
5. Steps (required) - Test steps (numbered steps separated by newlines)
6. Expected Result (required) - Expected outcome
7. Actual Result (optional) - Actual outcome (leave empty for new test cases)
8. Tab (optional) - Tab/section name (defaults to 'General')

Example CSV:
Title,Description,Priority,Type,Steps,Expected Result,Actual Result,Tab
Login with valid credentials,Test successful login,high,functional,"1. Open login page
2. Enter username
3. Enter password
4. Click login",User should be logged in,,Authentication

Missing columns in your file: """ + ', '.join(missing_columns)
            
            raise HTTPException(status_code=400, detail=format_guide)
        
        # Validate and import test cases
        imported_count = 0
        errors = []
        
        for index, row in df.iterrows():
            row_num = index + 2  # +2 because Excel rows start at 1 and header is row 1
            
            try:
                # Get values (handle both capitalized and lowercase column names)
                title = str(row.get('Title', row.get('title', ''))).strip()
                description = str(row.get('Description', row.get('description', ''))).strip()
                priority = str(row.get('Priority', row.get('priority', 'medium'))).strip().lower()
                test_type = str(row.get('Type', row.get('type', 'functional'))).strip().lower()
                steps = str(row.get('Steps', row.get('steps', ''))).strip()
                expected_result = str(row.get('Expected Result', row.get('expected_result', row.get('expected result', '')))).strip()
                actual_result = str(row.get('Actual Result', row.get('actual_result', row.get('actual result', '')))).strip()
                tab = str(row.get('Tab', row.get('tab', 'General'))).strip()
                
                # Skip empty rows
                if not title or title == 'nan':
                    continue
                
                # Validate priority
                valid_priorities = ['low', 'medium', 'high']
                if priority not in valid_priorities:
                    errors.append(f"Row {row_num}: Invalid priority '{priority}'. Must be one of: {', '.join(valid_priorities)}")
                    continue
                
                # Validate type
                valid_types = ['functional', 'negative', 'ui/ux', 'smoke', 'regression', 'api']
                if test_type not in valid_types:
                    errors.append(f"Row {row_num}: Invalid type '{test_type}'. Must be one of: {', '.join(valid_types)}")
                    continue
                
                # Create test case
                tc_data = TestCaseCreate(
                    project_id=project_id,
                    tab=tab,
                    title=title,
                    description=description,
                    priority=priority,
                    type=test_type,
                    steps=steps,
                    expected_result=expected_result,
                    actual_result=actual_result if actual_result and actual_result != 'nan' else ""
                )
                
                # Convert to TestCase model
                tc = TestCase(**tc_data.model_dump(), created_by=current_user['username'])
                
                # Prepare document for MongoDB
                doc = tc.model_dump()
                doc['created_at'] = doc['created_at'].isoformat()
                doc['updated_at'] = doc['updated_at'].isoformat()
                if doc.get('executed_at'):
                    doc['executed_at'] = doc['executed_at'].isoformat()
                
                # Insert into database
                await db.test_cases.insert_one(doc)
                
                # Add tab to project if it doesn't exist
                if tab and tab != 'General':
                    await db.projects.update_one(
                        {"id": project_id},
                        {"$addToSet": {"tabs": tab}}
                    )
                
                imported_count += 1
                
            except Exception as e:
                errors.append(f"Row {row_num}: {str(e)}")
        
        # Return result with any errors
        result = {"imported_count": imported_count}
        if errors:
            result["errors"] = errors[:10]  # Limit to first 10 errors
            result["total_errors"] = len(errors)
        
        return result
        
    except pd.errors.EmptyDataError:
        raise HTTPException(status_code=400, detail="The uploaded file is empty")
    except pd.errors.ParserError as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Import failed: {str(e)}")


# ============ EXPORT ROUTES (Viewer can access) ============

@api_router.get("/test-cases/export/excel/{project_id}")
async def export_excel(project_id: str, current_user: dict = Depends(get_current_user)):
    project = await check_project_permission(project_id, current_user, "viewer")
    
    test_cases = await db.test_cases.find({"project_id": project_id, "is_template": False}, {"_id": 0}).to_list(1000)
    
    tabs_dict = {}
    for tc in test_cases:
        tab = tc.get('tab', 'General')
        if tab not in tabs_dict:
            tabs_dict[tab] = []
        tabs_dict[tab].append(tc)
    
    for tab in tabs_dict:
        tabs_dict[tab].sort(key=lambda x: x.get('created_at', ''))
    
    wb = Workbook()
    wb.remove(wb.active)
    
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    header_alignment = Alignment(horizontal="center", vertical="center")
    data_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    
    global_counter = 1
    for tab_name in sorted(tabs_dict.keys()):
        safe_sheet_name = tab_name[:31].replace('/', '-').replace('\\', '-').replace('*', '').replace('?', '').replace('[', '').replace(']', '')
        ws = wb.create_sheet(title=safe_sheet_name)
        
        headers = ['TC ID', 'Title', 'Description', 'Priority', 'Type', 'Steps', 'Expected Result', 'Actual Result', 'Status', 'Assigned To', 'Executed At']
        ws.append(headers)
        
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_alignment
        
        for tc in tabs_dict[tab_name]:
            tc_id = f"TC{str(global_counter).zfill(3)}"
            global_counter += 1
            
            executed_at = ''
            if tc.get('executed_at'):
                if isinstance(tc['executed_at'], str):
                    executed_at = tc['executed_at'].split('T')[0]
                else:
                    executed_at = tc['executed_at'].strftime('%Y-%m-%d')
            
            row_data = [
                tc_id,
                tc.get('title', ''),
                tc.get('description', ''),
                tc.get('priority', ''),
                tc.get('type', ''),
                tc.get('steps', ''),
                tc.get('expected_result', ''),
                tc.get('actual_result', ''),
                tc.get('status', ''),
                tc.get('assigned_to', ''),
                executed_at
            ]
            ws.append(row_data)
            
            current_row = ws.max_row
            for cell in ws[current_row]:
                cell.alignment = data_alignment
        
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
            ws.row_dimensions[row[0].row].height = None
        
        ws.column_dimensions['A'].width = 10
        ws.column_dimensions['B'].width = 25
        ws.column_dimensions['C'].width = 30
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 15
        ws.column_dimensions['F'].width = 35
        ws.column_dimensions['G'].width = 35
        ws.column_dimensions['H'].width = 35
        ws.column_dimensions['I'].width = 12
        ws.column_dimensions['J'].width = 15
        ws.column_dimensions['K'].width = 15
    
    if not tabs_dict:
        ws = wb.create_sheet(title="General")
        headers = ['TC ID', 'Title', 'Description', 'Priority', 'Type', 'Steps', 'Expected Result', 'Actual Result', 'Status', 'Assigned To', 'Executed At']
        ws.append(headers)
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_alignment
    
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={project['name']}_test_cases.xlsx"}
    )

@api_router.get("/test-cases/export/docx/{project_id}")
async def export_docx(project_id: str, current_user: dict = Depends(get_current_user)):
    project = await check_project_permission(project_id, current_user, "viewer")
    
    test_cases = await db.test_cases.find({"project_id": project_id, "is_template": False}, {"_id": 0}).to_list(1000)
    test_cases.sort(key=lambda x: (x.get('tab', 'General'), x.get('created_at', '')))
    
    doc = Document()
    
    title = doc.add_heading(f"{project['name']} - Test Cases", 0)
    title.alignment = 1  # Center
    
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    headers = ['TC ID', 'Tab', 'Title', 'Description', 'Priority', 'Type', 'Steps', 'Expected Result', 'Actual Result']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    for idx, tc in enumerate(test_cases, start=1):
        tc_id = f"TC{str(idx).zfill(3)}"
        row_cells = table.add_row().cells
        row_cells[0].text = tc_id
        row_cells[1].text = tc.get('tab', 'General')
        row_cells[2].text = tc.get('title', '')
        row_cells[3].text = tc.get('description', '')
        row_cells[4].text = tc.get('priority', '')
        row_cells[5].text = tc.get('type', '')
        row_cells[6].text = tc.get('steps', '')
        row_cells[7].text = tc.get('expected_result', '')
        row_cells[8].text = tc.get('actual_result', '')
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={project['name']}_test_cases.docx"}
    )


app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
